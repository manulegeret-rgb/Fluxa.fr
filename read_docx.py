from docx import Document

doc = Document('SEO/3 points critiques SEO.docx')

print("=== PARAGRAPHES ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text}")

print("\n=== TABLEAUX ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- Tableau {i+1} ---")
    for row in table.rows:
        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
        if row_text.strip():
            print(row_text)

print("\n=== SECTIONS ===")
print(f"Nombre de sections: {len(doc.sections)}")
print(f"Nombre de paragraphes: {len(doc.paragraphs)}")
print(f"Nombre de tableaux: {len(doc.tables)}")
