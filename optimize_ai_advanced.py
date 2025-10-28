#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Optimisation SEO avancée avec IA pour les articles Fluxa
Système intelligent d'amélioration de contenu
"""

import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class FluxaSEOOptimizer:
    """Optimiseur SEO intelligent pour les articles Fluxa"""

    def __init__(self):
        self.keywords_principal = [
            'automatisation', 'workflow', 'artisan', 'artisans',
            'tpe', 'pme', 'indépendant', 'indépendants',
            'fluxa', 'n8n', 'gestion'
        ]

        self.keywords_secondaire = [
            'digitalisation', 'gagner du temps', 'outil',
            'automatiser', 'efficacité', 'productivité',
            'entreprise', 'professionnel'
        ]

        self.bad_ctas = [
            r'cliquez?\s+ici',
            r'essai\s+gratuit',
            r'découvrez\s+la\s+démo\s+gratuite?',
            r'essayez\s+maintenant',
            r'profitez\s+de',
            r'offre\s+spéciale',
            r'inscription\s+gratuite?',
            r'testez\s+gratuitement'
        ]

    def clean_marketing_ctas(self, text):
        """Supprime ou remplace les CTAs marketing"""
        original_text = text

        for pattern in self.bad_ctas:
            # Remplacer par des alternatives neutres
            text = re.sub(pattern, 'En savoir plus sur Fluxa', text, flags=re.IGNORECASE)

        return text, (text != original_text)

    def optimize_meta_description(self, meta):
        """Optimise la longueur et le contenu de la meta description"""
        if not meta:
            return None

        # Nettoyer le préfixe "Meta description :"
        meta = re.sub(r'^meta\s+description\s*:\s*', '', meta, flags=re.IGNORECASE).strip()

        # Si trop longue, couper intelligemment
        if len(meta) > 160:
            # Couper à la dernière phrase complète avant 160 caractères
            sentences = re.split(r'[.!?]\s+', meta)
            meta_optimized = ""

            for sentence in sentences:
                if len(meta_optimized + sentence) <= 155:
                    meta_optimized += sentence + ". "
                else:
                    break

            meta = meta_optimized.strip()

        # Si trop courte, ajouter un CTA neutre
        if len(meta) < 120:
            if not meta.endswith('.'):
                meta += '.'
            meta += " Guide complet pour artisans et TPE."

        return meta

    def enrich_lexical_field(self, text):
        """Enrichit le champ lexical SEO du texte"""
        improvements = []

        # Remplacements intelligents pour enrichir le vocabulaire
        replacements = {
            r'\boutil\b': 'solution',
            r'\bfaire\b': 'mettre en place',
            r'\bbon\b': 'efficace',
            r'\brapide\b': 'performant',
            r'\bsimple\b': 'accessible',
            r'\bfacile\b': 'intuitif'
        }

        for pattern, replacement in replacements.items():
            if re.search(pattern, text, re.IGNORECASE):
                # Ne remplacer qu'une occurrence sur deux pour garder la naturalité
                matches = list(re.finditer(pattern, text, re.IGNORECASE))
                if len(matches) > 1:
                    # Remplacer la deuxième occurrence
                    match = matches[1]
                    text = text[:match.start()] + replacement + text[match.end():]
                    improvements.append(f"Enrichissement lexical: '{match.group()}' → '{replacement}'")

        return text, improvements

    def optimize_heading(self, heading, level):
        """Optimise un titre (H1, H2, H3)"""
        # Supprimer les guillemets inutiles
        heading = heading.strip('"\'')

        # Ajouter des power words pour le SEO
        power_words = {
            'H1': ['Guide Complet', 'en 2025', 'pour Artisans et TPE'],
            'H2': ['Comment', 'Pourquoi', 'Les Clés', 'Étape par Étape'],
            'H3': ['Concrètement', 'En Pratique', 'Exemple']
        }

        # S'assurer que les titres sont percutants
        if level == 'H1':
            # Le H1 doit contenir le mot-clé principal et l'année
            if '2025' not in heading:
                heading = heading.rstrip(':. ') + ' en 2025'

        return heading

    def calculate_keyword_density(self, text):
        """Calcule la densité des mots-clés principaux"""
        text_lower = text.lower()
        word_count = len(text.split())

        densities = {}
        for keyword in self.keywords_principal:
            count = text_lower.count(keyword.lower())
            density = (count / word_count * 100) if word_count > 0 else 0
            densities[keyword] = {
                'count': count,
                'density': density,
                'status': 'OK' if count >= 3 else 'LOW'
            }

        return densities

    def add_internal_links(self, text):
        """Ajoute des liens internes vers Fluxa de manière naturelle"""
        # Identifier les endroits où ajouter un lien
        link_patterns = [
            (r'(automatisation\s+pour\s+artisans)', r'\1 via Fluxa.fr'),
            (r'(solution\s+de\s+gestion)', r'\1 comme Fluxa'),
            (r'(outil\s+d\'automatisation)', r'\1 tel que Fluxa')
        ]

        for pattern, replacement in link_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                text = re.sub(pattern, replacement, text, count=1, flags=re.IGNORECASE)
                break  # Un seul lien interne par paragraphe

        return text

    def optimize_article(self, article_data, filename):
        """Optimise un article complet"""
        print(f"\n{'='*60}")
        print(f"Optimisation de : {filename}")
        print(f"{'='*60}")

        optimized = {
            'filename': filename,
            'paragraphs': [],
            'optimizations': [],
            'keyword_density': {},
            'issues_fixed': []
        }

        full_text = article_data['full_text']

        # Analyser la densité de mots-clés AVANT optimisation
        densities_before = self.calculate_keyword_density(full_text)

        meta_found = False

        for i, para in enumerate(article_data['paragraphs']):
            text = para['text']
            style = para['style']

            # 1. Traiter les titres (H1, H2, H3)
            if 'Heading 1' in style:
                text = self.optimize_heading(text, 'H1')
                optimized['optimizations'].append(f"H1 optimisé: {text[:50]}...")

            elif 'Heading 2' in style:
                text = self.optimize_heading(text, 'H2')

            elif 'Heading 3' in style:
                text = self.optimize_heading(text, 'H3')

            # 2. Traiter la meta description
            if not meta_found and 'meta' in text.lower() and 'description' in text.lower():
                meta_optimized = self.optimize_meta_description(text)
                if meta_optimized:
                    text = f"Meta description : {meta_optimized}"
                    optimized['optimizations'].append(f"Meta optimisée: {len(meta_optimized)} caractères")
                    meta_found = True

            # 3. Supprimer les CTAs marketing
            text_cleaned, cta_removed = self.clean_marketing_ctas(text)
            if cta_removed:
                optimized['issues_fixed'].append(f"CTA marketing supprimé (ligne {i+1})")
                text = text_cleaned

            # 4. Enrichir le champ lexical
            text_enriched, improvements = self.enrich_lexical_field(text)
            if improvements:
                optimized['optimizations'].extend(improvements)
                text = text_enriched

            # 5. Ajouter des liens internes (max 1 par article)
            if i > 5 and i < 15 and len([opt for opt in optimized['optimizations'] if 'lien interne' in opt]) == 0:
                if any(kw in text.lower() for kw in ['automatisation', 'solution', 'outil']):
                    text_with_link = self.add_internal_links(text)
                    if text_with_link != text:
                        optimized['optimizations'].append("Lien interne ajouté vers Fluxa")
                        text = text_with_link

            optimized['paragraphs'].append({
                'text': text,
                'style': style
            })

        # Analyser la densité APRÈS optimisation
        full_text_optimized = '\n\n'.join([p['text'] for p in optimized['paragraphs']])
        densities_after = self.calculate_keyword_density(full_text_optimized)
        optimized['keyword_density'] = densities_after

        # Afficher le résumé
        print(f"OK - {len(optimized['optimizations'])} optimisations appliquees")
        print(f"OK - {len(optimized['issues_fixed'])} problemes corriges")

        # Afficher les mots-clés améliorés
        print("\nDensite des mots-cles (avant -> apres):")
        for kw in ['automatisation', 'workflow', 'artisan', 'tpe']:
            before = densities_before.get(kw, {}).get('count', 0)
            after = densities_after.get(kw, {}).get('count', 0)
            status = "OK" if after >= before else "="
            print(f"  {status} {kw}: {before} -> {after} occurrences")

        return optimized

    def save_to_docx(self, optimized_data, output_path):
        """Sauvegarde l'article optimisé en .docx"""
        doc = Document()

        # Configurer les styles de base
        for para_data in optimized_data['paragraphs']:
            text = para_data['text']
            style_name = para_data['style']

            # Ajouter le paragraphe
            if 'Heading 1' in style_name:
                p = doc.add_heading(text, level=1)
            elif 'Heading 2' in style_name:
                p = doc.add_heading(text, level=2)
            elif 'Heading 3' in style_name:
                p = doc.add_heading(text, level=3)
            else:
                p = doc.add_paragraph(text)

        # Sauvegarder
        doc.save(output_path)
        print(f"OK - Sauvegarde: {os.path.basename(output_path)}")

def main():
    print("="*80)
    print("OPTIMISATION SEO AVANCÉE AVEC IA - FLUXA")
    print("="*80)

    # Charger les articles
    print("\nChargement des articles...")
    with open('articles_extracted.json', 'r', encoding='utf-8') as f:
        articles = json.load(f)

    print(f"OK - {len(articles)} articles charges")

    # Créer le dossier de sortie
    output_dir = r"C:\Users\Utilisateur\Documents\Applications personnalisées\fluxa-artisans-automations-main\Articles optimisés SEO"
    os.makedirs(output_dir, exist_ok=True)

    # Initialiser l'optimiseur
    optimizer = FluxaSEOOptimizer()

    # Optimiser chaque article
    all_optimizations = {}

    for i, (filename, article_data) in enumerate(articles.items(), 1):
        print(f"\n[{i}/{len(articles)}]")

        # Optimiser l'article
        optimized = optimizer.optimize_article(article_data, filename)
        all_optimizations[filename] = optimized

        # Sauvegarder en .docx
        base_name = filename.replace('.docx', '')
        output_filename = f"{base_name}_SEO.docx"
        output_path = os.path.join(output_dir, output_filename)

        optimizer.save_to_docx(optimized, output_path)

        print(f"OK - Article {i}/{len(articles)} termine\n")

    # Sauvegarder le rapport JSON
    report_path = os.path.join(output_dir, "rapport_optimisations.json")
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(all_optimizations, f, ensure_ascii=False, indent=2)

    print("\n" + "="*80)
    print("OPTIMISATION TERMINEE")
    print("="*80)
    print(f"\nOK - {len(articles)} articles optimises")
    print(f"OK - Sauvegardes dans: {output_dir}")
    print(f"OK - Rapport detaille: {report_path}")

if __name__ == "__main__":
    main()
